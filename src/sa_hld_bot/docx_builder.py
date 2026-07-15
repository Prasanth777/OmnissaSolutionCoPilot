from __future__ import annotations

import re
from copy import deepcopy
from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor

from .catalog import Product


# Brand palette modeled on the reference design document (navy 122B4A system).
NAVY_HEX = "122B4A"
NAVY = RGBColor(0x12, 0x2B, 0x4A)
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK = RGBColor(0x0B, 0x25, 0x45)
BODY_GREY = RGBColor(0x40, 0x40, 0x40)
MUTED = RGBColor(0x5B, 0x66, 0x73)
TABLE_HEADER = NAVY_HEX
BAND_FILL = "F2F5FA"
LIGHT_FILL = "F4F6F9"
BORDER_HEX = "C9D3E0"
WHITE = "FFFFFF"
CONTENT_WIDTH_DXA = 9360
HEADING_FONT = "Calibri Light"
BODY_FONT = "Calibri"
HEADING_NUM_ID = "199"


def _clean_text(text: object, max_len: int = 1200) -> str:
    raw = str(text or "").strip()
    if not raw:
        return ""
    cleaned = re.sub(r"(^|\s)[#*_`>-]+", " ", raw)
    cleaned = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    return cleaned[:max_len]


def _is_unknown(value: object) -> bool:
    text = _clean_text(value).lower()
    return not text or "unknown" in text or "to be confirmed" in text or text == "tbd"


def _answer(answers: dict[str, str], key: str, default: str = "To be confirmed") -> str:
    value = _clean_text(answers.get(key, ""), 800)
    return default if _is_unknown(value) else value


def _diagram_header(row: dict[str, str], idx: int, max_len: int = 180) -> str:
    for key in ("figure_caption", "caption", "slide_title", "title"):
        value = _clean_text(row.get(key, ""), max_len)
        if value:
            value = re.sub(r"^figure\s+\d+\s*[:.\-]?\s*", "", value, flags=re.IGNORECASE)
            return value or f"Architecture Diagram {idx}"
    return f"Architecture Diagram {idx}"


def _canon_url(url: str) -> str:
    return str(url or "").strip().rstrip("/")


def _row_text(row: dict[str, str]) -> str:
    return " ".join(
        _clean_text(row.get(key, ""), 300)
        for key in ("topic", "slide_title", "title", "caption", "figure_caption", "section_heading", "context_text", "embed_text", "page_url")
    ).lower()


def _figure_signature(row: dict[str, str]) -> str:
    text = _row_text(row)
    text = re.sub(r"\bfigure\s+\d+\b", " ", text)
    text = re.sub(r"\btable\s+\d+\b", " ", text)
    text = re.sub(r"[^a-z0-9]+", " ", text).strip()
    if text:
        return text[:120]
    return _canon_url(str(row.get("image_url") or row.get("local_path") or "")).lower()


def _load_balancer_applicable(answers: dict[str, str]) -> bool:
    values = " ".join(
        _clean_text(answers.get(key, ""), 300).lower()
        for key in ("load_balancer", "load_balancer_name", "load_balancer_placement", "horizon_access_topology", "horizon_external_access", "uag_arch_track")
    )
    return "load balancer" in values or "load balancing" in values or "load-balanced" in values or values.startswith("yes")


def _is_blocked_generic_connection(row: dict[str, str]) -> bool:
    text = _row_text(row)
    if "load balanc" in text:
        return False
    return ("internal connection" in text or "external connection" in text) and "blast" not in text


def _infer_diagram_site(row: dict[str, str]) -> str:
    explicit = str(row.get("site_topology") or "")
    text = f" {_row_text(row)} "
    if "cloud pod" in text or " cpa " in text:
        return "multi"
    if any(term in text for term in (
        "multi-site", "multisite", "multi site", "multi-datacentre", "multi-datacenter",
        "active-active", "active/active", "active active",
        "active-passive", "active/passive", "active passive",
        "stretched cluster", "stretched vsan", "vsan stretched",
        "preferred site", "secondary site", "witness site", "data site to data site",
        "site 1", "site 2", "site 3",
    )):
        return "multi"
    if explicit in {"multisite", "multisite_active_active", "multisite_active_passive", "cloud_pod"}:
        return "multi"
    if explicit == "single_site" or "single-site" in text or "single site" in text:
        return "single"
    return ""


def _figure_conflicts_with_answers(row: dict[str, str], answers: dict[str, str], selected_products: list[str] | None = None) -> bool:
    """Delegate to the shared image-selection rules so DOCX figures are filtered
    with exactly the same answer-aware logic as the PPT selection."""
    try:
        from .image_select import figure_conflicts

        return figure_conflicts(row, answers, selected_products)
    except Exception:
        return _legacy_figure_conflicts_with_answers(row, answers)


def _legacy_figure_conflicts_with_answers(row: dict[str, str], answers: dict[str, str]) -> bool:
    requested_site = _answer(answers, "site_topology").lower()
    availability = _answer(answers, "availability_requirements").lower()
    wants_single = "single" in requested_site and "multi" not in requested_site and "active" not in availability
    wants_multi = "multi" in requested_site or "active" in availability
    diagram_site = _infer_diagram_site(row)
    if (wants_single and diagram_site == "multi") or (wants_multi and diagram_site == "single"):
        return True

    requested_dmz_raw = _answer(answers, "horizon_dmz_design").lower()
    requested_dmz = (
        "double" if "double" in requested_dmz_raw
        else "single" if "single" in requested_dmz_raw
        else "none" if "no dmz" in requested_dmz_raw or "internal only" in requested_dmz_raw
        else ""
    )
    text = _row_text(row)
    diagram_dmz_raw = str(row.get("dmz_design", "") or "").lower()
    diagram_dmz = (
        "double" if "double" in diagram_dmz_raw or "double dmz" in text
        else "single" if "single" in diagram_dmz_raw or "single dmz" in text
        else ""
    )
    if requested_dmz and diagram_dmz and requested_dmz != diagram_dmz:
        return True
    if requested_dmz == "none" and diagram_dmz in {"single", "double"}:
        return True
    return False


def _figure_topic(row: dict[str, str]) -> str:
    return _clean_text(row.get("topic") or row.get("slide_title") or row.get("figure_caption") or row.get("caption"), 180)


def _source_name(row: dict[str, str]) -> str:
    title = _clean_text(row.get("title", ""), 140)
    if title:
        return title
    url = str(row.get("page_url") or "").rstrip("/").split("/")[-1].replace("-", " ")
    return url.title() if url else "the referenced Tech Zone guidance"


def _figure_context(row: dict[str, str]) -> str:
    title = _diagram_header(row, 0)
    section = _clean_text(row.get("section_heading", ""), 140)
    source = _source_name(row)
    if section and section.lower() not in title.lower():
        return f"The figure '{title}' from the {section} section of {source}"
    return f"The figure '{title}' from {source}"


def _component_list(row: dict[str, str]) -> str:
    raw = row.get("components_shown") or []
    if isinstance(raw, str):
        items = [part.strip().replace("_", " ") for part in re.split(r"[,;]", raw) if part.strip()]
    else:
        items = [str(part).strip().replace("_", " ") for part in raw if str(part).strip()]
    return ", ".join(dict.fromkeys(items[:6]))


def _figure_explanation(row: dict[str, str], answers: dict[str, str], section: str = "") -> tuple[str, list[str]]:
    """Explain a figure as a build statement for the customer's environment.

    Every branch describes how the depicted architecture will be implemented
    for this customer, grounded in the interview answers, rather than a
    generic description of the diagram.
    """
    text = _row_text(row)
    context = _figure_context(row)
    components = _component_list(row)

    customer = _answer(answers, "customer_name", "the customer")

    def known(key: str, fallback: str) -> str:
        value = _answer(answers, key, "")
        return value if value and value != "To be confirmed" else fallback

    access = known("access_type", "the agreed access model")
    site = known("site_topology", "the agreed site topology")
    hosting = known("hosting_strategy", "the target platform")
    concurrency = known("workload_concurrency", "the agreed user concurrency")
    primary_site = known("primary_site", "the primary datacenter")
    secondary_sites = known("secondary_sites", "the secondary site")
    availability = known("availability_requirements", "the agreed availability model")
    lb_name = known("load_balancer_name", "the customer load balancer")
    lb_place = known("load_balancer_placement", "the agreed placement")
    dmz = known("horizon_dmz_design", "the agreed DMZ design")
    mfa_required = known("mfa_required", "")
    mfa_provider = known("mfa_provider", "the customer MFA provider")
    identity = known("identity_source", "the customer directory")
    cert = known("cert_type", "the agreed SSL certificate")
    fqdn = known("fqdn_strategy", "the agreed FQDN approach")
    uag_nics = known("uag_nic_config", "the agreed NIC layout")
    hzc_provider = known("horizon_cloud_provider", "the selected cloud provider")
    hzc_connectivity = known("horizon_cloud_connectivity", "the agreed access path")
    dr_scenarios = known("dr_scenarios", "the agreed recovery scenarios")
    pod_model = known("horizon_pod_block_model", "a scaled Horizon pod")
    internal_in_scope = "internal" in access.lower() or "both" in access.lower()

    # --- Horizon Cloud Service / Horizon Edge -------------------------------
    if "horizon edge" in text or "control plane" in text or "horizon cloud" in text:
        if section == "networking" or any(k in text for k in ("vnet", "subnet", "vpc", "connectivity", "networking", "ports")):
            return (
                f"{context} maps to the network layout that will be built in {customer}'s {hzc_provider} environment.",
                [
                    f"Dedicated subnets will be provisioned in {customer}'s {hzc_provider} network for the Horizon "
                    "Edge management components, the Unified Access Gateways, and the desktop capacity, with routing "
                    "and network security rules scoped to each tier.",
                    "The Edge requires outbound TCP 443 to the Omnissa control plane and name resolution to the "
                    "control-plane endpoints; no inbound connectivity from the internet to the management subnet is "
                    "required.",
                    f"User access will follow the agreed connectivity model ({hzc_connectivity}), and the required "
                    "address space will be reserved with the network team before deployment.",
                ],
            )
        if any(k in text for k in ("preparing", "subscription", "provider requirements", "capacity provider")):
            return (
                f"{context} shows the {hzc_provider} groundwork that will be completed before {customer}'s Horizon "
                "Edge is deployed.",
                [
                    f"A dedicated {hzc_provider} subscription/account scope will be prepared with the required "
                    "quotas, resource groups, and service principals for the Horizon Edge deployment.",
                    "Networking prerequisites (address space, DNS, outbound internet path) and identity integration "
                    f"with {identity} will be validated as part of the readiness checklist.",
                    "Once prerequisites pass validation, the Edge deployment is driven from the Horizon Universal "
                    "Console with no customer-managed brokering infrastructure to build.",
                ],
            )
        if any(k in text for k in ("pool", "images", "entitle", "assignment")):
            return (
                f"{context} shows the workflow that will be used to publish desktops and applications to "
                f"{customer}'s users.",
                [
                    f"Golden images will be imported and maintained in the Horizon Universal Console, then assigned "
                    f"to pools and pool groups sized for {concurrency.lower()}.",
                    f"Entitlements will be granted to {identity} groups so access control stays aligned with the "
                    "customer's existing directory administration model.",
                    "Image lifecycle (patching, versioning, rollback) will be operated through the console and is "
                    "documented in the operations section of this design.",
                ],
            )
        if section == "security" or any(k in text for k in ("interaction", "security", "trust")):
            return (
                f"{context} shows the trust boundary between {customer}'s environment and the Omnissa-managed "
                "control plane.",
                [
                    "All communication is initiated outbound from the Horizon Edge over TLS (TCP 443); Omnissa has "
                    "no inbound network path into the customer environment.",
                    f"User credentials are validated against {identity} inside {customer}'s environment; the control "
                    "plane orchestrates brokering and management but does not hold desktop session traffic.",
                    "This split keeps session data and workloads under customer control while Omnissa operates the "
                    "management services, which simplifies the security review.",
                ],
            )
        return (
            f"{context} shows how the Horizon Edge and the Omnissa-managed Horizon Control Plane "
            f"will operate in {customer}'s environment.",
            [
                f"The Horizon Edge will be deployed into {customer}'s {hzc_provider} capacity and will connect "
                "outbound over TCP 443 to the Omnissa-hosted control plane; no inbound connectivity from the "
                "control plane into the customer environment is required.",
                f"Brokering, entitlement, image, and monitoring services will be consumed from the control plane, "
                f"while all user sessions and desktop workloads remain inside {customer}'s {hzc_provider} environment.",
                f"The Edge will be sized for {concurrency.lower()}, and user access will follow the agreed "
                f"connectivity model ({hzc_connectivity}).",
            ],
        )
    if any(k in text for k in ("active-active", "active passive", "active-passive", "multi-site", "multisite", "stretched", "secondary site", "witness site", "cloud pod", " cpa ")):
        return (
            f"{context} shows the multi-site availability pattern that will be implemented for {customer}.",
            [
                f"Horizon capacity will be distributed across {primary_site} and {secondary_sites} in a "
                f"{availability.lower()} model, with Cloud Pod Architecture providing global entitlements and "
                "inter-pod session brokering.",
                f"Under normal operation users will be brokered to their home site; on a site outage, sessions will "
                f"be re-brokered to the surviving site in line with the agreed recovery scenarios ({dr_scenarios}).",
                "Replication of golden images, application packages, and user profile data between the sites will be "
                "defined during detailed design so both sites can serve their assigned users independently.",
            ],
        )
    if "blast" in text or "blast" in str(row.get("page_url", "")).lower():
        first = (
            f"External users will authenticate through {lb_name} to the Unified Access Gateway appliances in the "
            f"{dmz}, and the Blast Extreme session will then be established through UAG "
            "(TCP/UDP 8443 or 443 at the edge, TCP 22443 to the Horizon Agent)."
        )
        if access.lower().startswith("internal"):
            first = (
                "Users will connect to the Connection Servers for authentication, and the Blast Extreme session "
                "will then flow directly from the Horizon Client to the assigned desktop or RDSH host on TCP 22443 "
                "(UDP 22443 optional)."
            )
        second = (
            f"Internal users will connect directly to the Connection Servers, with their Blast sessions flowing "
            "straight to the Horizon Agent, bypassing the DMZ path."
            if internal_in_scope and not access.lower().startswith("internal")
            else f"Only Blast Extreme will be enabled as the display protocol for {customer}; PCoIP and RDP paths shown in general guidance are out of scope."
        )
        return (
            f"{context} maps directly to the network path that will carry {customer}'s Horizon sessions.",
            [
                first,
                second,
                f"Firewall rules between the client, DMZ, management, and workload segments will be opened per "
                "Omnissa network-port guidance for Blast Extreme, and validated with the network team before go-live.",
            ],
        )
    if "load balanc" in text:
        return (
            f"{context} reflects the load-balancing tier that will be built for {customer} using {lb_name}.",
            [
                f"{lb_name} will present a single virtual service {lb_place.lower()}, with health monitors "
                "determining which appliances receive new sessions; users always connect to the service name, "
                "never to an individual node.",
                f"The virtual service will carry the published Horizon FQDN ({fqdn}) and the {cert}, so DNS records "
                "and certificate subject names must match before cutover.",
                "Persistence, timeout, and health-check settings will follow Omnissa load-balancing guidance and "
                "will be validated during implementation testing.",
            ],
        )
    if "uag" in text or "unified access gateway" in text or "dmz" in text:
        return (
            f"{context} shows the secure edge that will publish {customer}'s Horizon services to external users.",
            [
                f"Unified Access Gateway appliances will be deployed with {uag_nics} in a {dmz}, "
                f"terminating TLS with the {cert} and forwarding only authenticated Horizon traffic to the "
                "internal Connection Servers.",
                f"{lb_name} will distribute external sessions across the UAG appliances so the edge tier remains "
                "available during maintenance or an appliance failure.",
                "Only the required Horizon ports will be permitted across the outer and inner firewall boundaries, "
                "keeping the management and workload networks isolated from the DMZ.",
            ],
        )
    if "true sso" in text or "enrollment server" in text:
        return (
            f"{context} shows how {customer}'s users will reach a Windows session without a second credential prompt.",
            [
                f"After {mfa_provider} authentication, True SSO will issue a short-lived certificate through the "
                f"Enrollment Server and an enterprise certificate authority integrated with {identity}, completing "
                "the Windows logon silently.",
                "The Enrollment Server will be deployed alongside the Connection Servers, with a dedicated CA "
                "template scoped to True SSO issuance.",
                "Certificate authority ownership, template configuration, and key length will be confirmed with the "
                "customer PKI team during detailed design.",
            ],
        )
    if "authentication" in text or "saml" in text or "radius" in text or "mfa" in text or "pass-through" in text:
        mfa_clause = (
            f"with multi-factor authentication enforced by {mfa_provider}"
            if mfa_required.lower().startswith("yes")
            else "per the agreed authentication policy"
        )
        return (
            f"{context} corresponds to the authentication flow that will protect {customer}'s Horizon access.",
            [
                f"Users will be validated against {identity} {mfa_clause} before any desktop or application session "
                "is brokered.",
                "For external users the authentication decision will be taken at the Unified Access Gateway tier, "
                "so unauthenticated traffic never reaches the internal network.",
                "Session and conditional-access policies will follow the security baseline captured in the Security "
                "Standards section of this document.",
            ],
        )
    if any(k in text for k in ("pod and block", "scaled horizon pod", "block design", "pod design")):
        return (
            f"{context} translates to the pod and block layout planned for {customer}.",
            [
                f"{pod_model} will host the Horizon management components, with resource blocks added as repeatable "
                f"capacity units sized for {concurrency.lower()}.",
                f"The management block will run the Connection Servers, databases, and supporting services on the "
                f"{hosting.lower()} platform at {primary_site}, separated from desktop and RDSH workload clusters.",
                "Block boundaries, cluster sizing, and growth increments will be validated against the workload "
                "profile during detailed design.",
            ],
        )
    if "app volumes" in text or "app attach" in text or "package" in text:
        av_scope = known("app_volumes_scope", "the agreed application delivery use cases")
        av_storage = known("app_volumes_storage", "the agreed storage platform")
        return (
            f"{context} shows the application-delivery layer that will be built for {customer} with App Volumes.",
            [
                f"App Volumes will deliver {av_scope.lower()}, attaching application packages to user sessions at "
                "logon so the golden images stay generic and simple to maintain.",
                f"Packages will be stored on {av_storage.lower()}, with the App Volumes Managers deployed "
                "redundantly behind a load-balanced service endpoint.",
                "Package replication, entitlement model, and database availability will be finalized in the App "
                "Volumes detailed design section.",
            ],
        )
    if "dynamic environment manager" in text or " dem " in f" {text} " or "profile" in text or "fslogix" in text:
        profile_strategy = known("dem_profile_strategy", "the agreed profile strategy")
        file_shares = known("dem_file_shares", "highly available SMB shares")
        return (
            f"{context} shows the user-environment layer that will personalize {customer}'s sessions.",
            [
                f"User settings and profiles will follow {profile_strategy.lower()}, keeping the desktops stateless "
                "while users retain their personalization across sessions.",
                f"Configuration and profile data will be hosted on {file_shares.lower()}, which must be reachable "
                "from every desktop and RDSH host at logon.",
                "Share sizing, availability, and GPO integration will be confirmed in the detailed design and "
                "operational handover.",
            ],
        )
    if "logical" in text or "component" in text or "core components" in text or components:
        component_clause = f" The components shown ({components}) each map to a section of this document." if components else ""
        return (
            f"{context} shows the components that will make up {customer}'s Horizon platform.{component_clause}",
            [
                f"The platform will be deployed {hosting.lower()} with {access.lower()} access, sized for "
                f"{concurrency.lower()}.",
                f"Brokering will be provided by the Connection Servers, secure external access by Unified Access "
                f"Gateway, identity by {identity}, and desktop/application capacity by the workload blocks described "
                "in the detailed design.",
                "Any component shown but not confirmed during the design interview is recorded as an assumption or "
                "open item rather than committed scope.",
            ],
        )
    return (
        f"{context} supports the proposed design for {customer}.",
        [
            f"The architecture shown will be realized on the {hosting.lower()} platform with {site.lower()} topology "
            f"and {access.lower()} access.",
            f"Sizing and placement decisions will align to {concurrency.lower()} and the requirements captured in "
            "this document.",
            "Values not yet confirmed are tracked as assumptions or open items for detailed design.",
        ],
    )


class HldDocxBuilder:
    """Build a formal HLD DOCX modeled on the Peninsula-style design document.

    Formatting system: A4 pages, navy-branded numbered headings, an automatic
    table of contents plus lists of tables/figures/design decisions (Word
    fields), SEQ-numbered captions, branded tables, and running headers and
    footers with page numbering.
    """

    def __init__(self) -> None:
        self._used_image_paths: set[str] = set()
        self._used_image_signatures: set[str] = set()
        self._figure_no = 0
        self._table_no = 0
        self._decision_no = 0
        self._first_h1_written = False
        self._product_keys: list[str] = []
        self._used_content_keys: set[str] = set()
        self._used_explanation_keys: set[int] = set()
        self._current_section: str = ""

    def build(
        self,
        output_path: Path,
        customer_name: str,
        selected_products: list[Product],
        questionnaire: dict[str, str],
        rag_narrative: dict[str, str],
        references: list[str],
        image_rows: list[dict[str, str]],
        custom_sections: list[dict] | None = None,
        excluded_sections: set[str] | list[str] | None = None,
    ) -> Path:
        self._used_image_paths = set()
        self._used_image_signatures = set()
        self._figure_no = 0
        self._table_no = 0
        self._decision_no = 0
        self._first_h1_written = False
        self._product_keys = [product.key for product in selected_products]
        self._used_content_keys = set()
        self._used_explanation_keys = set()
        self._current_section = ""

        image_rows = [
            row for row in image_rows
            if row.get("image_type", "architecture_diagram") == "architecture_diagram"
            and Path(str(row.get("local_path", ""))).exists()
        ]
        used_refs = self._used_references(references, image_rows)

        doc = Document()
        self._configure_document(doc)
        self._cover(doc, customer_name, selected_products, questionnaire)

        # Front matter: TOC + lists of tables/figures/design decisions.
        doc.add_section(WD_SECTION.NEW_PAGE)
        self._set_running_header_footer(doc.sections[-1], customer_name, questionnaire)
        self._front_matter(doc)

        # Main body.
        doc.add_section(WD_SECTION.NEW_PAGE)
        excluded = {str(s).lower() for s in (excluded_sections or [])}
        products_in_scope = [p for p in selected_products if p.key not in excluded]
        if "key_contacts" not in excluded:
            self._key_contacts(doc, questionnaire)
        if "overview" not in excluded and "summary" not in excluded:
            self._overview(doc, customer_name, selected_products, questionnaire, rag_narrative)
        if "requirements" not in excluded:
            self._requirements(doc, questionnaire)
        if "solution_overview" not in excluded and "architecture" not in excluded:
            self._solution_overview(doc, questionnaire, rag_narrative, image_rows)
        self._detailed_design(doc, products_in_scope, questionnaire, rag_narrative, image_rows)
        if "networking" not in excluded:
            self._networking(doc, questionnaire, rag_narrative, image_rows)
        if "security" not in excluded:
            self._security(doc, questionnaire, rag_narrative, image_rows)
        if "business_continuity" not in excluded and "operations" not in excluded:
            self._business_continuity(doc, questionnaire, rag_narrative, image_rows)
        for section in custom_sections or []:
            self._custom_section(doc, section, questionnaire, image_rows)
        if "additional_views" not in excluded:
            self._additional_views(doc, image_rows, questionnaire)
        self._references(doc, used_refs)
        if "review_acceptance" not in excluded:
            self._review_acceptance(doc, questionnaire)

        self._highlight_to_be_confirmed(doc)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        return output_path

    # ------------------------------------------------------------------ setup

    def _configure_document(self, doc: Document) -> None:
        section = doc.sections[0]
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

        styles = doc.styles
        normal = styles["Normal"]
        normal.font.name = BODY_FONT
        normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        normal.font.size = Pt(11)
        normal.font.color.rgb = BODY_GREY
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.15

        for name, size, color, before, after in (
            ("Heading 1", 20, NAVY, 20, 10),
            ("Heading 2", 16, NAVY, 14, 7),
            ("Heading 3", 13, DARK_BLUE, 10, 5),
            ("Heading 4", 11.5, DARK_BLUE, 8, 4),
            ("Heading 5", 11, DARK_BLUE, 6, 3),
        ):
            style = styles[name]
            style.font.name = HEADING_FONT
            style._element.rPr.rFonts.set(qn("w:ascii"), HEADING_FONT)
            style._element.rPr.rFonts.set(qn("w:hAnsi"), HEADING_FONT)
            style.font.size = Pt(size)
            style.font.color.rgb = color
            style.font.bold = True
            style.font.italic = False
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)
            style.paragraph_format.keep_with_next = True

        self._ensure_paragraph_style(doc, "Caption Text", size=9.5, color=DARK_BLUE, bold=True, before=8, after=3)
        self._ensure_paragraph_style(doc, "Figure Caption", size=9.5, color=MUTED, italic=True, after=8)
        self._ensure_paragraph_style(doc, "Figure Explanation", size=10, color=INK, before=2, after=4)
        self._ensure_paragraph_style(doc, "Design Decision Text", size=10, color=INK, before=2, after=2)
        self._ensure_paragraph_style(doc, "Reference Link", size=9.5, color=DARK_BLUE, after=3)
        self._ensure_paragraph_style(
            doc, "Front Matter Heading", size=20, color=NAVY, bold=True, before=6, after=12, font=HEADING_FONT
        )

        self._enable_heading_numbering(doc)
        self._enable_update_fields_on_open(doc)

    def _enable_update_fields_on_open(self, doc: Document) -> None:
        """Ask Word to refresh TOC/SEQ/PAGE fields when the document opens."""
        try:
            settings = doc.settings.element
        except Exception:
            return
        update = settings.find(qn("w:updateFields"))
        if update is None:
            update = OxmlElement("w:updateFields")
            settings.append(update)
        update.set(qn("w:val"), "true")

    def _highlight_to_be_confirmed(self, doc: Document) -> None:
        """Highlight every generated ``To be confirmed`` marker in yellow."""
        roots = [doc.element.body]
        seen_parts: set[int] = set()
        for section in doc.sections:
            for container in (section.header, section.footer):
                part_id = id(container.part)
                if part_id not in seen_parts:
                    seen_parts.add(part_id)
                    roots.append(container._element)

        for root in roots:
            for run in list(root.iter(qn("w:r"))):
                text = "".join(node.text or "" for node in run.iter(qn("w:t")))
                if not re.search(r"to be confirmed", text, flags=re.IGNORECASE):
                    continue

                parent = run.getparent()
                if parent is None:
                    continue
                insert_at = parent.index(run)
                segments = re.split(r"(to be confirmed)", text, flags=re.IGNORECASE)
                for segment in (part for part in segments if part):
                    replacement = deepcopy(run)
                    for child in list(replacement):
                        if child.tag != qn("w:rPr"):
                            replacement.remove(child)

                    text_node = OxmlElement("w:t")
                    if segment[:1].isspace() or segment[-1:].isspace():
                        text_node.set(qn("xml:space"), "preserve")
                    text_node.text = segment
                    replacement.append(text_node)

                    if segment.lower() == "to be confirmed":
                        run_properties = replacement.get_or_add_rPr()
                        existing = run_properties.find(qn("w:highlight"))
                        if existing is not None:
                            run_properties.remove(existing)
                        highlight = OxmlElement("w:highlight")
                        highlight.set(qn("w:val"), "yellow")
                        run_properties.append(highlight)

                    parent.insert(insert_at, replacement)
                    insert_at += 1
                parent.remove(run)

    def _enable_heading_numbering(self, doc: Document) -> None:
        """Attach multilevel numbering (1, 1.1, 1.1.1 ...) to Heading 1-5."""
        try:
            from docx.opc.constants import RELATIONSHIP_TYPE as RT

            numbering = doc.part.part_related_by(RT.NUMBERING).element
        except Exception:
            return
        if numbering is None:
            return

        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), HEADING_NUM_ID)
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "multilevel")
        abstract.append(multi)
        for ilvl in range(5):
            lvl = OxmlElement("w:lvl")
            lvl.set(qn("w:ilvl"), str(ilvl))
            start = OxmlElement("w:start")
            start.set(qn("w:val"), "1")
            lvl.append(start)
            fmt = OxmlElement("w:numFmt")
            fmt.set(qn("w:val"), "decimal")
            lvl.append(fmt)
            pstyle = OxmlElement("w:pStyle")
            pstyle.set(qn("w:val"), f"Heading{ilvl + 1}")
            lvl.append(pstyle)
            suffix = OxmlElement("w:suff")
            suffix.set(qn("w:val"), "space")
            lvl.append(suffix)
            text = OxmlElement("w:lvlText")
            text.set(qn("w:val"), ".".join(f"%{i + 1}" for i in range(ilvl + 1)))
            lvl.append(text)
            jc = OxmlElement("w:lvlJc")
            jc.set(qn("w:val"), "left")
            lvl.append(jc)
            ppr = OxmlElement("w:pPr")
            ind = OxmlElement("w:ind")
            ind.set(qn("w:left"), "0")
            ind.set(qn("w:firstLine"), "0")
            ppr.append(ind)
            lvl.append(ppr)
            abstract.append(lvl)

        first_num = numbering.find(qn("w:num"))
        if first_num is not None:
            first_num.addprevious(abstract)
        else:
            numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), HEADING_NUM_ID)
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), HEADING_NUM_ID)
        num.append(abstract_ref)
        numbering.append(num)

        for level in range(1, 6):
            style = doc.styles[f"Heading {level}"]
            ppr = style._element.get_or_add_pPr()
            num_pr = OxmlElement("w:numPr")
            ilvl = OxmlElement("w:ilvl")
            ilvl.set(qn("w:val"), str(level - 1))
            num_pr.append(ilvl)
            num_id = OxmlElement("w:numId")
            num_id.set(qn("w:val"), HEADING_NUM_ID)
            num_pr.append(num_id)
            ppr.append(num_pr)

    # -------------------------------------------------------------- fields

    def _add_field(self, paragraph, instr: str, placeholder: str = "") -> None:
        """Insert a Word field (TOC, SEQ, PAGE...) with placeholder text."""
        run = paragraph.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr_el = OxmlElement("w:instrText")
        instr_el.set(qn("xml:space"), "preserve")
        instr_el.text = f" {instr} "
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        run._r.append(begin)
        run._r.append(instr_el)
        run._r.append(separate)
        if placeholder:
            text = OxmlElement("w:t")
            text.set(qn("xml:space"), "preserve")
            text.text = placeholder
            run._r.append(text)
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.append(end)

    # ------------------------------------------------------ header / footer

    def _set_running_header_footer(self, section, customer_name: str, answers: dict[str, str]) -> None:
        project = _answer(answers, "project_name", "Architecture Design")
        customer = customer_name or "Customer"

        section.header.is_linked_to_previous = False
        header = section.header.paragraphs[0]
        for run in list(header.runs):
            run._r.getparent().remove(run._r)
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = header.add_run(f"{customer}  |  {project}")
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED
        self._paragraph_bottom_border(header, NAVY_HEX, size=6)

        section.footer.is_linked_to_previous = False
        footer = section.footer.paragraphs[0]
        for run in list(footer.runs):
            run._r.getparent().remove(run._r)
        footer.paragraph_format.tab_stops.add_tab_stop(Inches(6.27), WD_ALIGN_PARAGRAPH.RIGHT)
        left = footer.add_run(f"{customer} - High-Level Design")
        left.font.size = Pt(9)
        left.font.color.rgb = MUTED
        footer.add_run("\t")
        page_lead = footer.add_run("Page ")
        page_lead.font.size = Pt(9)
        page_lead.font.color.rgb = MUTED
        self._add_field(footer, "PAGE", "1")
        of_run = footer.add_run(" of ")
        of_run.font.size = Pt(9)
        of_run.font.color.rgb = MUTED
        self._add_field(footer, "NUMPAGES", "1")
        for run in footer.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = MUTED

    # ---------------------------------------------------------------- cover

    def _cover(self, doc: Document, customer_name: str, products: list[Product], answers: dict[str, str]) -> None:
        customer = customer_name or "Customer"
        product_names = ", ".join(product.title for product in products)
        title_text = f"Omnissa {products[0].title} Detailed Design" if len(products) == 1 else "Omnissa Architecture Detailed Design"

        band = doc.add_paragraph()
        band.paragraph_format.space_after = Pt(120)
        self._paragraph_shading(band, NAVY_HEX)
        band_run = band.add_run("  HIGH-LEVEL DESIGN")
        band_run.font.name = HEADING_FONT
        band_run.font.size = Pt(12)
        band_run.font.bold = True
        band_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        title = doc.add_paragraph()
        title.paragraph_format.space_after = Pt(6)
        run = title.add_run(title_text)
        run.font.name = HEADING_FONT
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = NAVY

        subtitle = doc.add_paragraph()
        subtitle.paragraph_format.space_after = Pt(30)
        sub = subtitle.add_run(customer)
        sub.font.name = HEADING_FONT
        sub.font.size = Pt(18)
        sub.font.color.rgb = BLUE

        stamp = doc.add_paragraph()
        stamp.paragraph_format.space_after = Pt(2)
        stamp_run = stamp.add_run(date.today().strftime("%B %Y"))
        stamp_run.font.size = Pt(12)
        stamp_run.font.color.rgb = MUTED

        version = doc.add_paragraph()
        version.paragraph_format.space_after = Pt(36)
        version_run = version.add_run(f"DOCUMENT VERSION {_answer(answers, 'document_version', '0.1 - Draft').upper()}")
        version_run.font.size = Pt(10)
        version_run.font.bold = True
        version_run.font.color.rgb = MUTED

        self._rule(doc)
        rows = [
            ("Project / Service", _answer(answers, "project_name", "Omnissa EUC Architecture")),
            ("Products", product_names or "To be confirmed"),
            ("Prepared By", _answer(answers, "prepared_by")),
            ("Customer Contacts", _answer(answers, "customer_contacts")),
            ("Version", _answer(answers, "document_version", "0.1 - Draft")),
            ("Date", date.today().strftime("%B %d, %Y")),
        ]
        self._simple_table(doc, rows, widths=(2500, 6200), header=None)

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.add_run("Document purpose: ").bold = True
        p.add_run(
            "This high-level design describes the proposed Omnissa architecture, key requirements, design "
            "decisions, network and security considerations, and recovery assumptions. It is based on the "
            "customer inputs captured during the design interview and Omnissa Tech Zone reference guidance."
        )

    # --------------------------------------------------------- front matter

    def _front_matter(self, doc: Document) -> None:
        doc.add_paragraph("Table of Contents", style="Front Matter Heading")
        toc = doc.add_paragraph()
        self._add_field(
            toc,
            'TOC \\o "1-3" \\h \\z \\u',
            "The table of contents will populate when the document is opened in Microsoft Word.",
        )

        spacer = doc.add_paragraph()
        spacer.add_run().add_break(WD_BREAK.PAGE)

        doc.add_paragraph("List of Tables", style="Front Matter Heading")
        tables_field = doc.add_paragraph()
        self._add_field(tables_field, 'TOC \\h \\z \\c "Table"', "The list of tables will populate when opened in Word.")

        doc.add_paragraph("List of Figures", style="Front Matter Heading")
        figures_field = doc.add_paragraph()
        self._add_field(figures_field, 'TOC \\h \\z \\c "Figure"', "The list of figures will populate when opened in Word.")

        doc.add_paragraph("List of Design Decisions", style="Front Matter Heading")
        decisions_field = doc.add_paragraph()
        self._add_field(
            decisions_field,
            'TOC \\h \\z \\c "DesignDecision"',
            "The list of design decisions will populate when opened in Word.",
        )

    def _h1(self, doc: Document, text: str) -> None:
        heading = doc.add_heading(text, level=1)
        if self._first_h1_written:
            heading.paragraph_format.page_break_before = True
        self._first_h1_written = True

    # ------------------------------------------------------------- sections

    def _key_contacts(self, doc: Document, answers: dict[str, str]) -> None:
        self._h1(doc, "Key Contacts")
        rows = [
            ("Prepared by", _answer(answers, "prepared_by")),
            ("Customer contacts", _answer(answers, "customer_contacts")),
            ("Reviewers / approvers", _answer(answers, "reviewers")),
            ("Operations owner", _answer(answers, "operations_owner")),
        ]
        self._numbered_table(doc, "Key Contacts", ("Role", "Name / Responsibility"), rows)

    def _overview(self, doc: Document, customer_name: str, products: list[Product], answers: dict[str, str], narrative: dict[str, str]) -> None:
        self._h1(doc, "Overview")
        self._current_section = "overview"
        self._narrative(doc, narrative.get("summary", ""), f"This document describes the proposed Omnissa architecture for {customer_name or 'the customer'}.")
        self._paragraph(
            doc,
            "The design is structured to capture requirements, assumptions, constraints, solution architecture, detailed component design, "
            "networking, security standards, and recovery considerations."
        )

        doc.add_heading("Audience", level=2)
        for item in (
            "Project executive sponsor",
            "Desktop and EUC operations leads",
            "Application operations leads",
            "Cloud, infrastructure, network, and security architects",
            "Implementation engineers responsible for detailed design and deployment",
        ):
            self._bullet(doc, item)

        doc.add_heading("Document Reference", level=2)
        rows = [
            ("Customer", customer_name or "Customer"),
            ("Industry", _answer(answers, "industry")),
            ("Products", ", ".join(product.title for product in products) or "To be confirmed"),
            ("Primary objective", _answer(answers, "project_scope")),
            ("Document status", _answer(answers, "document_version", "0.1 - Draft")),
        ]
        self._numbered_table(doc, "Document Reference", ("Field", "Value"), rows)

    def _requirements(self, doc: Document, answers: dict[str, str]) -> None:
        self._h1(doc, "Requirements and Considerations")
        self._paragraph(doc, "This section summarizes the business requirements, technical requirements, constraints, and risks that shape the HLD.")

        doc.add_heading("Business Requirements", level=2)
        self._numbered_table(
            doc,
            "Business Requirements",
            ("Requirement", "Design Input"),
            [
                ("Business drivers", _answer(answers, "business_drivers")),
                ("Primary objective", _answer(answers, "project_scope")),
                ("User personas", _answer(answers, "users_personas")),
                ("Success criteria", _answer(answers, "success_criteria")),
                ("In scope", _answer(answers, "in_scope")),
                ("Out of scope", _answer(answers, "out_of_scope")),
            ],
        )

        doc.add_heading("Technical Requirements", level=2)
        self._numbered_table(
            doc,
            "Technical Requirements",
            ("Requirement", "Design Input"),
            [
                ("Workloads / delivery model", _answer(answers, "horizon_use_cases", _answer(answers, "project_scope"))),
                ("Expected scale / concurrency", _answer(answers, "workload_concurrency")),
                ("Hosting strategy", _answer(answers, "hosting_strategy")),
                ("Site topology", _answer(answers, "site_topology")),
                ("Identity source", _answer(answers, "identity_source")),
                ("MFA", f"{_answer(answers, 'mfa_required')} - {_answer(answers, 'mfa_provider')}"),
                ("External access", _answer(answers, "access_type")),
                ("Load balancer", _answer(answers, "load_balancer")),
                ("Load balancer name / platform", _answer(answers, "load_balancer_name")),
                ("Load balancer placement", _answer(answers, "load_balancer_placement")),
                ("FQDN / DNS", _answer(answers, "fqdn_strategy")),
                ("Certificate", _answer(answers, "cert_type")),
            ],
        )

        doc.add_heading("Constraints", level=2)
        self._numbered_table(
            doc,
            "Constraints",
            ("Constraint", "Impact"),
            [
                ("Project constraints", _answer(answers, "constraints")),
                ("Network services", _answer(answers, "dns_dhcp_ntp")),
                ("Network segments", _answer(answers, "network_segments")),
                ("Firewall and ports", _answer(answers, "firewall_ports")),
                ("Open items", _answer(answers, "open_items")),
            ],
        )

        doc.add_heading("Risks", level=2)
        self._numbered_table(
            doc,
            "Risks",
            ("Risk", "Mitigation / Note"),
            [
                ("Known risks", _answer(answers, "risks")),
                ("Assumptions", _answer(answers, "assumptions")),
                ("Certificate ownership", _answer(answers, "certificate_owner")),
                ("Operational ownership", _answer(answers, "operations_owner")),
            ],
        )

    def _solution_overview(self, doc: Document, answers: dict[str, str], narrative: dict[str, str], images: list[dict[str, str]]) -> None:
        self._h1(doc, "Solution Overview")
        self._current_section = "solution"
        self._narrative(doc, narrative.get("architecture", ""), "The solution architecture will be finalized from the confirmed requirements and Tech Zone design guidance.")
        self._decision(doc, "Hosting and topology", f"{_answer(answers, 'hosting_strategy')} deployment with {_answer(answers, 'site_topology')} topology.")
        self._decision(doc, "Access model", f"{_answer(answers, 'access_type')} with {_answer(answers, 'load_balancer')} load balancing posture.")
        self._add_figures(doc, images, ("overall", "high level", "logical components", "logical architecture"), limit=2, answers=answers)

    def _detailed_design(self, doc: Document, products: list[Product], answers: dict[str, str], narrative: dict[str, str], images: list[dict[str, str]]) -> None:
        for product in products:
            self._h1(doc, f"{product.title} Detailed Design")
            self._current_section = "design"
            self._narrative(doc, narrative.get(product.key, ""), product.summary)
            if product.key == "horizon_8":
                self._horizon_design(doc, answers, images)
            elif product.key == "app_volumes":
                self._app_volumes_design(doc, answers, images)
            elif product.key == "dynamic_environment_manager":
                self._dem_design(doc, answers, images)
            elif product.key == "unified_access_gateway":
                self._uag_design(doc, answers, images)
            else:
                self._product_design_table(doc, product, answers)
                self._add_figures(doc, images, (product.title.lower(), product.key.replace("_", " ")), limit=1, answers=answers)

    def _horizon_design(self, doc: Document, answers: dict[str, str], images: list[dict[str, str]]) -> None:
        doc.add_heading("Horizon 8 Architecture", level=2)
        self._numbered_table(
            doc,
            "Horizon Component Design",
            ("Component", "Design Input"),
            [
                ("Pod and block model", _answer(answers, "horizon_pod_block_model")),
                ("Image families", _answer(answers, "horizon_pool_model")),
                ("Desktop image version", _answer(answers, "desktop_image_version")),
                ("Server image version", _answer(answers, "server_image_version")),
                ("Connection Servers", _answer(answers, "horizon_connection_server_count")),
                ("External access", _answer(answers, "horizon_external_access", _answer(answers, "access_type"))),
                ("Access topology", _answer(answers, "horizon_access_topology")),
                ("Load balancer placement", _answer(answers, "load_balancer_placement")),
                ("DMZ design", _answer(answers, "horizon_dmz_design")),
                ("Display protocols", _answer(answers, "horizon_protocol_scope", "Blast Extreme only")),
                ("Event database", _answer(answers, "horizon_database_events")),
                ("Golden image strategy", _answer(answers, "horizon_golden_image")),
            ],
        )
        self._decision(doc, "Horizon access topology", f"{_answer(answers, 'horizon_external_access', _answer(answers, 'horizon_access_topology'))} using {_answer(answers, 'load_balancer_placement', _answer(answers, 'horizon_dmz_design'))} edge placement.")
        self._add_figures(doc, images, ("pod and block", "block design", "scaled horizon pod", "cloud pod"), limit=1, answers=answers)
        self._add_figures(doc, images, ("connection server", "horizon logical", "single-site", "multi-site"), limit=2, answers=answers)

    def _app_volumes_design(self, doc: Document, answers: dict[str, str], images: list[dict[str, str]]) -> None:
        self._numbered_table(
            doc,
            "App Volumes Design",
            ("Design Area", "Design Input"),
            [
                ("Use case", _answer(answers, "app_volumes_scope")),
                ("Architecture track", _answer(answers, "app_volumes_arch_track")),
                ("Design focus", _answer(answers, "app_volumes_design_focus")),
                ("Storage", _answer(answers, "app_volumes_storage")),
                ("Database", _answer(answers, "app_volumes_database")),
            ],
        )
        self._add_figures(doc, images, ("app volumes", "storage group", "apps on demand", "package", "database"), limit=2, answers=answers)

    def _dem_design(self, doc: Document, answers: dict[str, str], images: list[dict[str, str]]) -> None:
        self._numbered_table(
            doc,
            "Dynamic Environment Manager Design",
            ("Design Area", "Design Input"),
            [
                ("Management scope", _answer(answers, "dem_scope")),
                ("Architecture track", _answer(answers, "dem_arch_track")),
                ("Design focus", _answer(answers, "dem_design_focus")),
                ("File shares", _answer(answers, "dem_file_shares")),
                ("Profile strategy", _answer(answers, "dem_profile_strategy")),
            ],
        )
        self._add_figures(doc, images, ("dynamic environment manager", "dem", "profile", "configuration share", "fslogix"), limit=2, answers=answers)

    def _uag_design(self, doc: Document, answers: dict[str, str], images: list[dict[str, str]]) -> None:
        self._numbered_table(
            doc,
            "Unified Access Gateway Design",
            ("Design Area", "Design Input"),
            [
                ("NIC configuration", _answer(answers, "uag_nic_config")),
                ("Published services", _answer(answers, "uag_services")),
                ("Edge pattern", _answer(answers, "uag_edge_pattern")),
                ("Architecture track", _answer(answers, "uag_arch_track")),
                ("Design focus", _answer(answers, "uag_design_focus")),
            ],
        )
        self._add_figures(doc, images, ("unified access gateway", "uag", "dmz", "pass-through", "load balancing"), limit=2, answers=answers)

    def _product_design_table(self, doc: Document, product: Product, answers: dict[str, str]) -> None:
        prefix = product.key
        rows = [
            ("Architecture track", _answer(answers, f"{prefix}_arch_track")),
            ("Design focus", _answer(answers, f"{prefix}_design_focus")),
        ]
        self._numbered_table(doc, f"{product.title} Design Inputs", ("Design Area", "Design Input"), rows)

    def _networking(self, doc: Document, answers: dict[str, str], narrative: dict[str, str], images: list[dict[str, str]]) -> None:
        self._h1(doc, "Networking Requirements")
        self._current_section = "networking"
        self._narrative(
            doc,
            narrative.get("networking", ""),
            "Network requirements must be validated against the customer firewall, DNS, DHCP, NTP, and load-balancing standards.",
        )
        self._numbered_table(
            doc,
            "Network Requirements",
            ("Area", "Requirement / Assumption"),
            [
                ("Primary site", _answer(answers, "primary_site")),
                ("Secondary sites", _answer(answers, "secondary_sites", "Not applicable")),
                ("Access type", _answer(answers, "access_type")),
                ("Network segments", _answer(answers, "network_segments")),
                ("DNS / DHCP / NTP", _answer(answers, "dns_dhcp_ntp")),
                ("Firewall / ports", _answer(answers, "firewall_ports")),
                ("Load balancing", _answer(answers, "load_balancer")),
                ("Load balancer name / platform", _answer(answers, "load_balancer_name")),
                ("Load balancer placement", _answer(answers, "load_balancer_placement")),
                ("FQDN strategy", _answer(answers, "fqdn_strategy")),
            ],
        )
        self._add_figures(doc, images, ("blast", "blast extreme", "display protocol"), limit=1, answers=answers)
        if _load_balancer_applicable(answers):
            self._add_figures(doc, images, ("load balancing", "load-balanced", "load balancer", "connection server"), limit=1, answers=answers)
            if _answer(answers, "access_type").lower() != "internal users only":
                self._add_figures(doc, images, ("uag", "unified access gateway", "dmz"), limit=1, answers=answers)

    def _security(self, doc: Document, answers: dict[str, str], narrative: dict[str, str], images: list[dict[str, str]]) -> None:
        self._h1(doc, "Security Standards")
        self._current_section = "security"
        self._narrative(doc, narrative.get("security", ""), "Security standards should align to the customer identity, MFA, certificate, RBAC, logging, and hardening requirements.")
        self._numbered_table(
            doc,
            "Security Standards",
            ("Security Area", "Design Input"),
            [
                ("Security baseline", _answer(answers, "security_requirements")),
                ("Identity source", _answer(answers, "identity_source")),
                ("MFA", f"{_answer(answers, 'mfa_required')} - {_answer(answers, 'mfa_provider')}"),
                ("Certificate type", _answer(answers, "cert_type")),
                ("Certificate owner", _answer(answers, "certificate_owner")),
                ("RBAC", _answer(answers, "rbac_model")),
                ("Antivirus / hardening", _answer(answers, "antivirus_hardening")),
                ("Monitoring / logging", _answer(answers, "monitoring_logging")),
            ],
        )
        self._decision(doc, "Security posture", f"{_answer(answers, 'security_requirements')} with {_answer(answers, 'mfa_required')} MFA requirement.")
        self._add_figures(doc, images, ("authentication", "true sso", "access", "pass-through", "security"), limit=2, answers=answers)

    def _business_continuity(self, doc: Document, answers: dict[str, str], narrative: dict[str, str], images: list[dict[str, str]]) -> None:
        self._h1(doc, "Business Continuity and Recovery")
        self._current_section = "continuity"
        self._narrative(doc, narrative.get("operations", ""), "Availability and recovery design should be validated through detailed design and operational readiness workshops.")
        self._numbered_table(
            doc,
            "Disaster Recovery Scenarios",
            ("Scenario", "Design Response"),
            [
                ("Availability target", _answer(answers, "availability_requirements")),
                ("Backup expectations", _answer(answers, "backup_requirements")),
                ("DR scenarios", _answer(answers, "dr_scenarios")),
                ("Operations owner", _answer(answers, "operations_owner")),
                ("Monitoring and logging", _answer(answers, "monitoring_logging")),
                ("Open recovery items", _answer(answers, "open_items")),
            ],
        )
        if "multi-site" in _answer(answers, "site_topology").lower() or "active" in _answer(answers, "availability_requirements").lower():
            self._add_figures(doc, images, ("active-active", "active-passive", "multi-site", "cloud pod", "operations dashboard"), limit=2, answers=answers)

    def _custom_section(self, doc: Document, section: dict, answers: dict[str, str], images: list[dict[str, str]]) -> None:
        """Render a user-requested section (title + grounded narrative + figures)."""
        title = _clean_text(section.get("title", ""), 120) or "Additional Design Topic"
        self._h1(doc, title)
        self._current_section = "design"
        self._narrative(doc, section.get("content", ""), "Content for this section is pending detailed design.")
        keywords = tuple(k for k in (section.get("keywords") or []) if k)
        if keywords:
            self._add_figures(doc, images, keywords, limit=2, answers=answers)

    def _additional_views(self, doc: Document, images: list[dict[str, str]], answers: dict[str, str]) -> None:
        """Place any selected diagrams the section keyword filters did not use.

        The selection engine already validated every row against the answers, so
        dropping leftovers silently loses relevant content; instead they are
        documented here as supporting views.
        """
        remaining = [
            row for row in images
            if str(row.get("local_path", "")) not in self._used_image_paths
            and _figure_signature(row) not in self._used_image_signatures
            and not _is_blocked_generic_connection(row)
            and not _figure_conflicts_with_answers(row, answers, self._product_keys)
        ]
        if not remaining:
            return
        self._h1(doc, "Additional Architecture Views")
        self._current_section = "design"
        self._paragraph(
            doc,
            "The following supporting diagrams were matched to the design inputs and provide additional "
            "architectural detail for implementation planning.",
        )
        for row in remaining:
            self._add_figure(doc, row, answers)

    def _references(self, doc: Document, references: list[str]) -> None:
        self._h1(doc, "References")
        if not references:
            self._paragraph(doc, "No external source links were used in the generated content.")
            return
        for idx, ref in enumerate(references, start=1):
            p = doc.add_paragraph(style="Reference Link")
            p.add_run(f"{idx}. ")
            self._add_hyperlink(p, ref, ref)

    def _review_acceptance(self, doc: Document, answers: dict[str, str]) -> None:
        self._h1(doc, "Review and Acceptance")
        self._paragraph(doc, "The following sign-off table is provided for review tracking and acceptance of the high-level design.")
        self._table_no += 1
        self._table_caption(doc, self._table_no, "Review and Acceptance")
        rows = [
            (_answer(answers, "reviewers"), "Design review / approval", "", ""),
            (_answer(answers, "operations_owner"), "Operational acceptance", "", ""),
            (_answer(answers, "customer_contacts"), "Customer stakeholder acknowledgement", "", ""),
        ]
        self._generic_table(
            doc,
            headers=("Name / Team", "Responsibility", "Signature", "Date"),
            rows=rows,
            widths=(2800, 3300, 1800, 1300),
        )

    # -------------------------------------------------------------- figures

    def _add_figures(self, doc: Document, images: list[dict[str, str]], keywords: Iterable[str], limit: int = 1, answers: dict[str, str] | None = None) -> None:
        matched: list[dict[str, str]] = []
        kws = tuple(k.lower() for k in keywords)
        for row in images:
            local_path = str(row.get("local_path", ""))
            sig = _figure_signature(row)
            if not local_path or local_path in self._used_image_paths or sig in self._used_image_signatures:
                continue
            if _is_blocked_generic_connection(row):
                continue
            if _figure_conflicts_with_answers(row, answers or {}, getattr(self, "_product_keys", None)):
                continue
            text = _row_text(row)
            if any(keyword in text for keyword in kws):
                matched.append(row)
            if len(matched) >= limit:
                break

        for row in matched[:limit]:
            self._add_figure(doc, row, answers or {})

    def _add_figure(self, doc: Document, row: dict[str, str], answers: dict[str, str]) -> None:
        image_path = Path(str(row.get("local_path", "")))
        if not image_path.exists():
            return
        # Skip images whose file content duplicates an already-placed figure
        # (same diagram republished under a different URL/caption).
        try:
            from .image_select import image_content_keys

            content_keys = {k for k in image_content_keys(str(image_path)) if k}
        except Exception:
            content_keys = set()
        used_content = getattr(self, "_used_content_keys", None)
        if used_content is None:
            used_content = set()
            self._used_content_keys = used_content
        if content_keys & used_content:
            return
        used_content.update(content_keys)
        self._used_image_paths.add(str(image_path))
        self._used_image_signatures.add(_figure_signature(row))
        self._figure_no += 1
        title = _diagram_header(row, self._figure_no)
        doc.add_paragraph()
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(image_path), width=Inches(6.2))
        except Exception:
            return
        caption = doc.add_paragraph(style="Caption Text")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.add_run("Figure ")
        self._add_field(caption, "SEQ Figure \\* ARABIC", str(self._figure_no))
        caption.add_run(f": {title}")
        source = _clean_text(row.get("page_url", ""), 300)
        if source:
            src = doc.add_paragraph(style="Figure Caption")
            src.alignment = WD_ALIGN_PARAGRAPH.CENTER
            src.add_run("Source: ")
            self._add_hyperlink(src, source, source)
        lead, bullets = _figure_explanation(row, answers, self._current_section)
        # Never repeat an identical explanation elsewhere in the document: when a
        # second figure resolves to the same text, describe it as a complementary
        # view anchored on its own source section instead.
        explanation_key = hash(tuple(bullets))
        if explanation_key in self._used_explanation_keys:
            focus = _clean_text(row.get("section_heading", ""), 120) or title
            lead = (
                f"{_figure_context(row)} provides an additional view of the same design area, "
                f"focused on {focus}."
            )
            bullets = [
                f"The build decisions described for the earlier figures apply unchanged; this view highlights "
                f"{focus} in more detail.",
                "Use this view during implementation planning to validate component placement, dependencies, and "
                "firewall boundaries for the affected tier.",
            ]
        self._used_explanation_keys.add(explanation_key)
        explanation = doc.add_paragraph(style="Figure Explanation")
        explanation.paragraph_format.left_indent = Inches(0.18)
        explanation.paragraph_format.right_indent = Inches(0.18)
        self._paragraph_shading(explanation, BAND_FILL)
        explanation.add_run("Architecture explanation: ").bold = True
        explanation.add_run(lead)
        for bullet in bullets:
            self._bullet(doc, bullet)

    # ------------------------------------------------------ design decision

    def _decision(self, doc: Document, title: str, text: str) -> None:
        self._decision_no += 1
        caption = doc.add_paragraph(style="Caption Text")
        caption.add_run("Design Decision ")
        self._add_field(caption, "SEQ DesignDecision \\* ARABIC", str(self._decision_no))
        caption.add_run(f": {title}")

        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        self._set_table_width(table, (CONTENT_WIDTH_DXA,))
        self._set_table_borders(table)
        cell = table.rows[0].cells[0]
        self._cell_shading(cell, BAND_FILL)
        self._cell_left_accent(cell, NAVY_HEX)
        p = cell.paragraphs[0]
        p.style = doc.styles["Design Decision Text"]
        run = p.add_run("Decision: ")
        run.bold = True
        p.add_run(_clean_text(text, 500))
        doc.add_paragraph()

    # ---------------------------------------------------------------- tables

    def _table_caption(self, doc: Document, number: int, title: str) -> None:
        cap = doc.add_paragraph(style="Caption Text")
        cap.add_run("Table ")
        self._add_field(cap, "SEQ Table \\* ARABIC", str(number))
        cap.add_run(f": {title}")

    def _numbered_table(self, doc: Document, title: str, headers: tuple[str, str], rows: list[tuple[str, str]]) -> None:
        self._table_no += 1
        self._table_caption(doc, self._table_no, title)
        self._simple_table(doc, rows, widths=(3000, 6200), header=headers)

    def _simple_table(
        self,
        doc: Document,
        rows: list[tuple[str, str]],
        widths: tuple[int, int],
        header: tuple[str, str] | None,
    ) -> None:
        table = doc.add_table(rows=1 if header else 0, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        self._set_table_width(table, widths)
        self._set_table_borders(table)
        if header:
            hdr_row = table.rows[0]
            self._mark_header_row(hdr_row)
            for idx, text in enumerate(header):
                self._set_cell(hdr_row.cells[idx], text, bold=True, fill=TABLE_HEADER, color=RGBColor(0xFF, 0xFF, 0xFF))
        for row_idx, (label, value) in enumerate(rows):
            cells = table.add_row().cells
            band = BAND_FILL if row_idx % 2 else WHITE
            self._set_cell(cells[0], label, bold=True, fill=band)
            self._set_cell(cells[1], value or "To be confirmed", fill=band)
        doc.add_paragraph()

    def _generic_table(
        self,
        doc: Document,
        headers: tuple[str, ...],
        rows: list[tuple[str, ...]],
        widths: tuple[int, ...],
    ) -> None:
        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        self._set_table_width(table, widths)
        self._set_table_borders(table)
        hdr_row = table.rows[0]
        self._mark_header_row(hdr_row)
        for idx, text in enumerate(headers):
            self._set_cell(hdr_row.cells[idx], text, bold=True, fill=TABLE_HEADER, color=RGBColor(0xFF, 0xFF, 0xFF))
        for row_idx, row_values in enumerate(rows):
            cells = table.add_row().cells
            band = BAND_FILL if row_idx % 2 else WHITE
            for idx, value in enumerate(row_values[: len(headers)]):
                self._set_cell(cells[idx], value, bold=idx == 0, fill=band)
        doc.add_paragraph()

    def _mark_header_row(self, row) -> None:
        """Repeat the header row when the table breaks across pages."""
        tr_pr = row._tr.get_or_add_trPr()
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)

    def _set_table_borders(self, table) -> None:
        tbl_pr = table._tbl.tblPr
        borders = tbl_pr.find(qn("w:tblBorders"))
        if borders is None:
            borders = OxmlElement("w:tblBorders")
            tbl_pr.append(borders)
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            element = borders.find(qn(f"w:{edge}"))
            if element is None:
                element = OxmlElement(f"w:{edge}")
                borders.append(element)
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "4")
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), BORDER_HEX)

    def _set_table_width(self, table, widths: tuple[int, ...]) -> None:
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        tbl_w = tbl_pr.find(qn("w:tblW"))
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.append(tbl_w)
        tbl_w.set(qn("w:w"), str(sum(widths)))
        tbl_w.set(qn("w:type"), "dxa")
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                if idx >= len(widths):
                    continue
                cell.width = Pt(widths[idx] / 20)
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_w = tc_pr.find(qn("w:tcW"))
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:w"), str(widths[idx]))
                tc_w.set(qn("w:type"), "dxa")

    def _set_cell(self, cell, text: str, bold: bool = False, fill: str | None = None, color: RGBColor | None = None) -> None:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if fill:
            self._cell_shading(cell, fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(_clean_text(text, 900))
        run.bold = bold
        run.font.size = Pt(10)
        if color is not None:
            run.font.color.rgb = color

    # ---------------------------------------------------------------- prose

    def _paragraph(self, doc: Document, text: str) -> None:
        doc.add_paragraph(_clean_text(text, 1600) or "To be confirmed.")

    def _narrative(self, doc: Document, text: str, default: str = "") -> None:
        """Render narrative text as readable multi-sentence paragraphs.

        Strips leading inline labels (e.g. 'Executive Summary') left over from
        markdown cleaning, and splits long narrative blobs into paragraphs of
        about three sentences so sections don't read as a single wall of text.
        """
        cleaned = _clean_text(text, 2400) or _clean_text(default, 1600)
        if not cleaned:
            doc.add_paragraph("To be confirmed.")
            return
        cleaned = re.sub(
            r"^(executive summary|summary|overview|introduction)\s*[:\-]?\s*",
            "", cleaned, flags=re.IGNORECASE,
        )
        # If length truncation cut the text mid-sentence, drop the trailing fragment.
        if cleaned and cleaned[-1] not in ".!?":
            last_end = max(cleaned.rfind("."), cleaned.rfind("!"), cleaned.rfind("?"))
            if last_end > len(cleaned) // 2:
                cleaned = cleaned[: last_end + 1]
        sentences = re.split(r"(?<=[.!?])\s+", cleaned)
        chunk: list[str] = []
        for sentence in sentences:
            chunk.append(sentence)
            if len(chunk) == 3:
                doc.add_paragraph(" ".join(chunk))
                chunk = []
        if chunk:
            doc.add_paragraph(" ".join(chunk))

    def _bullet(self, doc: Document, text: str) -> None:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(_clean_text(text, 400))

    def _rule(self, doc: Document) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
        self._paragraph_bottom_border(p, NAVY_HEX, size=8)

    def _paragraph_bottom_border(self, paragraph, color_hex: str, size: int = 8) -> None:
        p_pr = paragraph._p.get_or_add_pPr()
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), str(size))
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), color_hex)
        p_bdr.append(bottom)

    def _ensure_paragraph_style(
        self,
        doc: Document,
        name: str,
        size: float,
        color: RGBColor,
        bold: bool = False,
        italic: bool = False,
        before: float = 0,
        after: float = 0,
        font: str = BODY_FONT,
    ) -> None:
        styles = doc.styles
        try:
            style = styles[name]
        except KeyError:
            style = styles.add_style(name, 1)
        style.font.name = font
        style._element.rPr.rFonts.set(qn("w:ascii"), font)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), font)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = bold
        style.font.italic = italic
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    def _cell_shading(self, cell, fill: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tc_pr.append(shd)
        shd.set(qn("w:fill"), fill)

    def _cell_left_accent(self, cell, color_hex: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = tc_pr.find(qn("w:tcBorders"))
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tc_pr.append(borders)
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "24")
        left.set(qn("w:space"), "0")
        left.set(qn("w:color"), color_hex)
        borders.append(left)

    def _paragraph_shading(self, paragraph, fill: str) -> None:
        p_pr = paragraph._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        p_pr.append(shd)

    def _add_hyperlink(self, paragraph, text: str, url: str) -> None:
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        run = OxmlElement("w:r")
        r_pr = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "1F4D78")
        r_pr.append(color)
        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        r_pr.append(underline)
        run.append(r_pr)
        text_element = OxmlElement("w:t")
        text_element.text = text
        run.append(text_element)
        hyperlink.append(run)
        paragraph._p.append(hyperlink)

    def _used_references(self, references: list[str], image_rows: list[dict[str, str]]) -> list[str]:
        refs: list[str] = []
        for ref in references or []:
            if ref:
                refs.append(_canon_url(ref))
        for row in image_rows:
            ref = _canon_url(str(row.get("page_url", "")))
            if ref:
                refs.append(ref)
        return list(dict.fromkeys(refs))
